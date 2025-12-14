from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import numpy as np # Importamos numpy por si acaso, aunque no se usa para generar el doc

def create_report():
    doc = Document()
    
    # --- Estilos del Título ---
    title = doc.add_heading('Informe Técnico: Desglose de Código MOEA/D', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Este documento detalla línea por línea la implementación del algoritmo MOEA/D aplicado al problema ZDT3, explicando la lógica matemática y funcional detrás de cada instrucción.')
    
    # --- Datos del Análisis (Tu código desglosado) ---
    # Formato: (Bloque de Código, Explicación Técnica)
    
    analysis_data = [
        ("class ZDT3:\n    def __init__(self, n_vars=30):", 
         "Define la clase del problema ZDT3. Por defecto, se establece en 30 variables de decisión (n_vars=30), tal como exige la competición."),
        
        ("self.bound_min = np.zeros(n_vars)\nself.bound_max = np.ones(n_vars)", 
         "Define el Espacio de Búsqueda. Todas las variables en ZDT3 deben estar normalizadas entre 0 y 1."),
        
        ("def evaluate(self, x):\n    x = np.clip(x, self.bound_min, self.bound_max)", 
         "Función de evaluación. Primero asegura (clip) que las variables no excedan los límites [0,1], corrigiendo posibles violaciones de los operadores evolutivos."),
        
        ("f1 = x[0]", 
         "Calcula el primer objetivo f1(x). En ZDT3, f1 depende solo de la primera variable."),
        
        ("g = 1 + 9 * np.mean(x[1:])", 
         "Calcula la función auxiliar 'g'. Involucra la media de todas las variables excepto la primera (x2 hasta x30)."),
        
        ("h = 1 - np.sqrt(f1 / g) - (f1 / g) * np.sin(10 * np.pi * f1)", 
         "Calcula la función auxiliar 'h'. Introduce la multimodalidad (múltiples óptimos locales) mediante el término seno (sin). Esto hace que el frente de Pareto sea discontinuo."),
        
        ("f2 = g * h\nreturn np.array([f1, f2])", 
         "Calcula el segundo objetivo f2. Retorna el vector de fitness [f1, f2]."),
        
        ("class MOEAD:\n    def __init__(...):", 
         "Constructor de la clase MOEA/D. Recibe el problema, tamaño de población (N), tamaño de vecindad (T) y generaciones."),
        
        ("self.CR = 0.5\nself.F = 0.5", 
         "Parámetros de Evolución Diferencial (DE). F es el factor de escala de mutación y CR es la probabilidad de cruce."),
        
        ("def initialize(self):", 
         "Método de inicialización. Se ejecuta una vez al principio para configurar pesos, vecinos y población inicial."),
        
        ("self.weights[i, 0] = i / (self.N - 1)\nself.weights[i, 1] = 1.0 - ...", 
         "Generación de Vectores de Peso. Crea una distribución uniforme de pesos (Lambda). Si N=100, crea vectores como [0,1], [0.01, 0.99]... hasta [1,0]."),
        
        ("self.weights[self.weights == 0] = 0.0001", 
         "Evita divisiones por cero en cálculos futuros sustituyendo valores 0 exactos por un valor muy pequeño."),
        
        ("dists[i, j] = np.linalg.norm(self.weights[i] - self.weights[j])", 
         "Calcula la Distancia Euclidiana entre todos los vectores de peso para determinar cuáles están cerca entre sí."),
        
        ("self.neighborhood = np.argsort(dists, axis=1)[:, :self.T]", 
         "Define la Vecindad B(i). Para cada subproblema, selecciona los 'T' índices de los vectores de peso más cercanos."),
        
        ("self.fitness_pop = np.array([...])\nself.z_ideal = np.min(...)", 
         "Evalúa la población inicial aleatoria e inicializa el Punto de Referencia Z*. Este vector contiene el mejor valor encontrado hasta ahora para f1 y f2."),
        
        ("def tchebycheff(self, fitness, weight_idx):", 
         "Función de Agregación (Escalarización). Convierte el problema multiobjetivo en uno de un solo objetivo para poder comparar soluciones."),
        
        ("diff = np.abs(fitness - self.z_ideal)\nreturn np.max(self.weights[weight_idx] * diff)", 
         "Implementa la fórmula de Tchebycheff: max{ lambda_i * |fi(x) - zi*| }. Busca minimizar la distancia ponderada al punto ideal."),
        
        ("def evolution_operator_DE(self, idx_subproblem):", 
         "Operador Evolutivo Completo. Combina Selección de Vecinos, Mutación DE, Cruce Binomial y Mutación Gaussiana."),
        
        ("neighbors = self.neighborhood[idx_subproblem]\nr1, r2, r3 = np.random.choice(...)", 
         "Selección de Padres (Mating). Elige aleatoriamente 3 índices (r1, r2, r3) EXCLUSIVAMENTE de la vecindad del subproblema actual."),
        
        ("mutant = x_r1 + self.F * (x_r2 - x_r3)", 
         "Mutación Diferencial (DE/rand/1). Crea un vector mutante basado en la diferencia de dos vecinos escalada por F."),
        
        ("crossover_mask = np.random.rand(n_vars) <= self.CR", 
         "Cruce Binomial. Genera una máscara booleana. Cada gen tiene una probabilidad CR de ser tomado del vector mutante."),
        
        ("offspring = np.where(crossover_mask, mutant, x_current)", 
         "Recombinación. Construye el hijo (offspring). Si la máscara es True toma del mutante, si es False conserva el valor del padre actual (x_current)."),
        
        ("if np.random.rand() < pr:\n    sigma = ...\n    offspring[j] += np.random.normal(0, sigma)", 
         "Mutación Gaussiana Polinomial. Con probabilidad 1/p, añade ruido gaussiano a una variable. Esto es crucial para escapar de óptimos locales y mantener diversidad."),
        
        ("def run(self):", 
         "Bucle principal del algoritmo evolutivo."),
        
        ("offspring_x = self.evolution_operator_DE(i)", 
         "Paso de Reproducción. Genera una nueva solución candidata 'y' para el subproblema 'i'."),
        
        ("self.z_ideal = np.min(np.vstack((self.z_ideal, offspring_fit)), axis=0)", 
         "Actualización de Z*. Comprueba si la nueva solución mejora el mejor valor histórico de f1 o f2 y actualiza el punto de referencia."),
        
        ("for j in neighbors:", 
         "Bucle de Actualización de Vecinos. El hijo generado no solo compite por el subproblema 'i', sino por todos los subproblemas en su vecindad."),
        
        ("if g_te_offspring <= g_te_neighbor:", 
         "Comparación Tchebycheff. Comprueba si la nueva solución 'y' es mejor que la solución actual 'x^j' PARA el vector de peso lambda^j."),
        
        ("self.population[j] = offspring_x", 
         "Reemplazo. Si el hijo es mejor, sobrescribe al individuo anterior en la población. (IMPORTANTE: Se recomienda añadir un límite de reemplazos para mantener la diversidad).")
    ]

    # --- Generación de la Tabla ---
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Encabezados de tabla
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fragmento de Código'
    hdr_cells[1].text = 'Análisis Funcional y Teórico'
    
    # Dar formato negrita y color de fondo a los encabezados
    for cell in hdr_cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.bold = True

    # Llenar la tabla
    for code, explanation in analysis_data:
        row_cells = table.add_row().cells
        
        # Columna Código
        p_code = row_cells[0].paragraphs[0]
        run_code = p_code.add_run(code)
        run_code.font.name = 'Courier New'
        run_code.font.size = Pt(9)
        run_code.font.color.rgb = RGBColor(0, 51, 102) # Azul oscuro
        
        # Columna Explicación
        p_expl = row_cells[1].paragraphs[0]
        p_expl.add_run(explanation)

    # --- Guardar ---
    filename = 'Informe_Codigo_MOEAD.docx'
    doc.save(filename)
    print(f"¡Listo! Se ha generado el archivo '{filename}' con el análisis detallado.")

if __name__ == "__main__":
    create_report()