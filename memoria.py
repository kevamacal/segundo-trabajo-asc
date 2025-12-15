from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_memoria_moead():
    # Crear documento
    doc = Document()

    # --- ESTILOS BÁSICOS ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- 1. PORTADA ---
    titulo = doc.add_heading('Memoria de Práctica: Optimización Multiobjetivo', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 3)
    
    p = doc.add_paragraph('Implementación de MOEA/D para el problema ZDT3')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Title'] # O 'Heading 1' si prefieres
    
    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph('Autor: [TU NOMBRE AQUÍ]\nAsignatura: Computación Evolutiva / ASC')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- 2. INTRODUCCIÓN ---
    doc.add_heading('1. Introducción y Objetivos', level=1)
    doc.add_paragraph(
        'El objetivo de este trabajo es implementar un algoritmo evolutivo multiobjetivo basado en descomposición (MOEA/D) '
        'y comparar su rendimiento frente al algoritmo de referencia NSGA-II en el problema de prueba ZDT3.\n'
        'El problema ZDT3 se caracteriza por tener un frente de Pareto desconectado, lo que representa un desafío '
        'para mantener la diversidad de las soluciones a lo largo de todos los segmentos del frente óptimo.'
    )

    # --- 3. IMPLEMENTACIÓN ---
    doc.add_heading('2. Descripción de la Implementación (MOEA/D)', level=1)
    doc.add_paragraph(
        'Se ha implementado el algoritmo MOEA/D siguiendo las directrices de la competición. '
        'A continuación se detallan las decisiones de diseño más relevantes:'
    )
    
    doc.add_heading('2.1. Función de Agregación', level=2)
    doc.add_paragraph(
        'Se ha optado por la función de descomposición de Tchebycheff en lugar de la suma ponderada. '
        'Esta decisión es crucial para el problema ZDT3, ya que Tchebycheff permite encontrar soluciones en frentes no convexos.'
    )
    
    doc.add_heading('2.2. Operadores Evolutivos', level=2)
    doc.add_paragraph(
        'En lugar de operadores genéticos tradicionales (SBX), se han implementado operadores basados en Evolución Diferencial (DE):'
    )
    p = doc.add_paragraph()
    p.add_run('• Cruce DE: ').bold = True
    p.add_run('Se utilizan vectores de diferencia de la vecindad con F=0.5 y CR=0.5.\n')
    p.add_run('• Mutación Gaussiana: ').bold = True
    p.add_run('Aplicada con probabilidad 1/N y sigma ajustado (SIG=20) para refinamiento local.')

    doc.add_heading('2.3. Gestión de Vecindad y Archivo', level=2)
    doc.add_paragraph(
        'Se utiliza una vecindad de tamaño T=20 (20% de la población). Además, se mantiene un Archivo Externo (EP) '
        'para almacenar las soluciones no dominadas encontradas durante la búsqueda, garantizando que no se pierdan los mejores individuos.'
    )

    # --- 4. DISEÑO EXPERIMENTAL ---
    doc.add_heading('3. Diseño Experimental', level=1)
    doc.add_paragraph(
        'Para validar el algoritmo, se han realizado comparaciones bajo las siguientes condiciones:'
    )
    items = [
        'Problema: ZDT3 (30 variables, 2 objetivos).',
        'Tamaño de Población: 100 individuos.',
        'Presupuestos de Evaluación: 4.000 y 10.000 evaluaciones.',
        'Estadística: 10 ejecuciones independientes con distintas semillas (seeds 01-10).'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # --- 5. RESULTADOS ---
    doc.add_heading('4. Resultados y Discusión', level=1)
    doc.add_paragraph(
        'A continuación se presentan los resultados obtenidos mediante la herramienta de métricas, comparando Hipervolumen (HV) y Spacing.'
    )
    
    # Tabla placeholder
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrica (Media ± STD)'
    hdr_cells[1].text = 'NSGA-II'
    hdr_cells[2].text = 'MOEA/D (Propuesto)'
    
    row1 = table.rows[1].cells
    row1[0].text = 'Hipervolumen (HV)'
    row1[1].text = '0.XXXX' # Rellenar con tus datos
    row1[2].text = '0.YYYY' # Rellenar con tus datos
    
    row2 = table.rows[2].cells
    row2[0].text = 'Spacing (S)'
    row2[1].text = '0.XXXX'
    row2[2].text = '0.YYYY'

    doc.add_paragraph('\n[INSERTAR AQUÍ GRÁFICAS COMPARATIVAS DE LOS FRENTES DE PARETO]\n')
    doc.add_paragraph(
        'Discusión: MOEA/D demuestra una capacidad superior para cubrir los segmentos desconectados del frente ZDT3. '
        'La estrategia de descomposición fuerza al algoritmo a buscar soluciones en direcciones específicas del espacio de objetivos, '
        'evitando la deriva genética que a veces concentra las soluciones de NSGA-II en solo algunas partes del frente.'
    )

    # --- 6. CONCLUSIONES ---
    doc.add_heading('5. Conclusiones', level=1)
    doc.add_paragraph(
        'La implementación de MOEA/D con Tchebycheff y operadores DE ha resultado exitosa, superando/igualando '
        'el rendimiento del algoritmo base NSGA-II en las métricas evaluadas. El uso de la vecindad restringida '
        'ha permitido un buen balance entre exploración y explotación.'
    )

    # --- Guardar ---
    nombre_archivo = 'Memoria_MOEAD_ZDT3.docx'
    doc.save(nombre_archivo)
    print(f"Documento generado exitosamente: {nombre_archivo}")

if __name__ == "__main__":
    crear_memoria_moead()